"# -*- coding: utf-8 -*-"
import nltk
# import goslate
import pymongo
import time
# from textblob import TextBlob
# from nltk.tokenize import TabTokenizer, SpaceTokenizer, BlanklineTokenizer
import fitz  # this is pymupdf
import re
# from reportlab.pdfgen import canvas
# from reportlab.lib.pagesizes import letter, A4
# from reportlab.lib.units import inch, cm
from reportlab import rl_config
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from bijoy2unicode import converter
# from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER, TA_JUSTIFY
import bn_unicode_converter.stm as engine
import bijoy_unicode_mapping.bangla_handler as translate
from docx import Document
from docx.shared import Pt

pdfmetrics.registerFont(TTFont('AmarBangla', r'D:\PythonProject\dictionary_project\fonts\AmarBangla.ttf'))
pdfmetrics.registerFont(TTFont('AmarBanglaBold', r'D:\PythonProject\dictionary_project\fonts\AmarBanglaBold.ttf'))
pdfmetrics.registerFont(TTFont('KalpanaANSI', r'D:\PythonProject\dictionary_project\fonts\Kalpana ANSI.ttf'))
pdfmetrics.registerFont(TTFont('KalpanaUNICODE', r'D:\PythonProject\dictionary_project\fonts\Kalpana UNICODE.ttf'))
pdfmetrics.registerFont(TTFont('EkusheySumit', r'D:\PythonProject\dictionary_project\fonts\Sumit_03-09-2005.ttf'))
pdfmetrics.registerFont(TTFont('HindSiliguri', r'D:\PythonProject\dictionary_project\fonts\HindSiliguri-Light.ttf'))
pdfmetrics.registerFont(TTFont('NotoSansBengali', r'D:\PythonProject\dictionary_project\fonts\NotoSansBengali-Light.ttf'))
pdfmetrics.registerFont(TTFont('ShonarBangla', r'D:\PythonProject\dictionary_project\fonts\Shonar.ttf'))
pdfmetrics.registerFont(TTFont('SutonnyMJ', r'D:\PythonProject\dictionary_project\fonts\bndictionary_app\typeface.ttf'))
pdfmetrics.registerFont(TTFont('NIRMALA', r'D:\PythonProject\dictionary_project\fonts\NIRMALA.TTF'))
pdfmetrics.registerFont(TTFont('SolaimanLipi', r'D:\PythonProject\dictionary_project\fonts\SolaimanLipi.ttf'))
pdfmetrics.registerFont(TTFont('Galada', r'D:\PythonProject\dictionary_project\fonts\Galada-Regular.ttf'))
pdfmetrics.registerFont(TTFont('TNR', r'D:\PythonProject\dictionary_project\fonts\times-new-roman.ttf'))
pdfmetrics.registerFont(TTFont('Kalpurush', r'D:\PythonProject\dictionary_project\fonts\Kalpurush.ttf'))
pdfmetrics.registerFont(TTFont('Mina', r'D:\PythonProject\dictionary_project\fonts\Mina-Bold.ttf'))
pdfmetrics.registerFont(TTFont('NotoSerifBengali', r'D:\PythonProject\dictionary_project\fonts\bongcalender_app\unidefault.ttf'))
pdfmetrics.registerFont(TTFont('SiyamRupaliRoboto', r'D:\PythonProject\dictionary_project\fonts\boitoi_app\SiyamRupaliRoboto.ttf'))
pdfmetrics.registerFont(TTFont('Nikosh', r'D:\PythonProject\dictionary_project\fonts\boitoi_app\bn\Nikosh.ttf'))
pdfmetrics.registerFont(TTFont('Mbong', r'D:\PythonProject\dictionary_project\fonts\bongcalender_app\mbong.ttf'))
pdfmetrics.registerFont(TTFont('RobotoThin', r'D:\PythonProject\dictionary_project\fonts\bongcalender_app\Roboto-Thin.ttf'))
pdfmetrics.registerFont(TTFont('Bangla', r'D:\PythonProject\dictionary_project\fonts\boitoi_app\bn\Bangla.ttf'))

rl_config._SAVED['canvas_basefontname'] = 'ShonarBangla'
# rl_config._startUp()

pdf = SimpleDocTemplate("sample.pdf")

styles = getSampleStyleSheet()
styles.add(ParagraphStyle(fontName='NIRMALA', name='NIRMALA', leading=15, fontSize=10, justifyBreaks=1))
styles.add(ParagraphStyle(fontName='AmarBangla', name='AmarBangla', leading=15, fontSize=12))
styles.add(ParagraphStyle(fontName='ShonarBangla', name='ShonarBangla', leading=15, fontSize=12))
styles.add(ParagraphStyle(fontName='SutonnyMJ', name='SutonnyMJ', leading=15, fontSize=12))
styles.add(ParagraphStyle(fontName='KalpanaANSI', name='Kalpana', leading=15, fontSize=12))
styles.add(ParagraphStyle(fontName='EkusheySumit', name='Ekushey', leading=15, fontSize=12))
styles.add(ParagraphStyle(fontName='HindSiliguri', name='HindSiliguri', leading=15, fontSize=12))
styles.add(ParagraphStyle(fontName='NotoSansBengali', name='NotoSansBengali', leading=15, fontSize=12))
styles.add(ParagraphStyle(fontName='SolaimanLipi', name='SolaimanLipi', leading=15, fontSize=12))
styles.add(ParagraphStyle(fontName='Galada', name='Galada', leading=15, fontSize=12))
styles.add(ParagraphStyle(fontName='TNR', name='TNR', leading=15, fontSize=12))
styles.add(ParagraphStyle(fontName='Kalpurush', name='Kalpurush', leading=15, fontSize=12))
styles.add(ParagraphStyle(fontName='Mina', name='Mina', leading=15, fontSize=12))
styles.add(ParagraphStyle(fontName='NotoSerifBengali', name='NotoSerifBengali', leading=15, fontSize=12))
styles.add(ParagraphStyle(fontName='SiyamRupaliRoboto', name='SiyamRupaliRoboto', leading=15, fontSize=12))
styles.add(ParagraphStyle(fontName='Nikosh', name='Nikosh', leading=15, fontSize=12))
styles.add(ParagraphStyle(fontName='Mbong', name='Mbong', leading=15, fontSize=12))
styles.add(ParagraphStyle(fontName='RobotoThin', name='RobotoThin', leading=15, fontSize=12))
styles.add(ParagraphStyle(fontName='Bangla', name='Bangla', leading=15, fontSize=12))

stopwords = nltk.corpus.stopwords.words('english') # take the list of enlish stopword(like the, a, an, he, she, his, him etc. )


# test = converter.Unicode()
# another_text = "বেষ্টিত"
# toPrint=test.convertUnicodeToBijoy(another_text)
# print(toPrint)

def remove_stopwords(txt_tokenized):
    txt_clean =[word for word in txt_tokenized if word not in stopwords]
    return txt_clean

# filename = r'D:\PythonProject\dictionary_project\Malgudi Days.pdf'
# filename2 = r'D:\PythonProject\dictionary_project\I Too Had a Love Story.pdf'
filename3 = r'D:\PythonProject\dictionary_project\Dharmayoddha Kalki _ Avatar of Vishnu.pdf'
# filename4 = r'D:\PythonProject\dictionary_project\Digital Electronics And Logic Design.pdf'
client = pymongo.MongoClient("mongodb://localhost:27017/")
mydb = client["bengali_dictionary_app"]
mycol = mydb["words"]


doc = fitz.open(filename3)     # or fitz.Document(filename)


build_flow_obj = []
# page = doc[61]
sstr = """<!DOCTYPE html><html>
            <head>
                <meta charset="UTF-8" />
            </head>
            <body>
        """
flow_obj = Paragraph(sstr)
build_flow_obj.append(flow_obj)

starttime = time.time()

print("অনুগ্রহ".encode("utf-8"))
print(b'\xe0\xa6\x85\xe0\xa6\xa8\xe0\xa7\x81\xe0\xa6\x97\xe0\xa7\x8d\xe0\xa6\xb0\xe0\xa6\xb9'.decode("utf-8"))

# gs = goslate.Goslate()
# gs.translate(word, 'bn') 
test = converter.Unicode()


# document = Document()
# style = document.styles['Normal']
# font = style.font
# font.name = 'AmarBangla'
# font.size = Pt(12)


for page in doc:
    # print(page.getText())
    if page.number == 20:
        break
    print(f"Page: {page.number}")
    text = page.getText()
    lowerText = text.lower()
    words = []
    words_meanings = []
    flow_object = []
    # print(stopwords[0:10])
    data = re.split("[\s.,!?:;'\"-]+",lowerText)
    # print(words)
    cln_data = remove_stopwords(data)
    words = list(dict.fromkeys(cln_data)) #remove duplicate wor from the list
    # print(words)

    flow_object.append(text)

    for word in words:
        wr = mycol.find({"word": word.upper()})
        try:
            bn_meaning = [w["bn_meanings"] for w in wr][0]
            d = word, bn_meaning
            # print(re.UNICODE(bn_meaning))
            # print(bn_meaning)
            words_meanings.append(d)
        except IndexError:
            pass
        except TypeError:
            pass

    build_text="<br/>"
    for tag in flow_object:
        build_text+=tag
    # print(build_text)
    flow_obj = Paragraph(build_text)
    build_flow_obj.append(flow_obj)

    build_text="<br/><br/>"
    for meanings in words_meanings:
        wrd = meanings[0]
        bn_m = meanings[1]
        # print(u"অনুগ্রহ")
        print(bn_m)
        # tst = bn_m.encode("utf-8")
        # print(tst)
        # bn_m_en = bn_m.encode("utf-8").decode('unicode-escape')
        # bn_m_en = [ bn_m[c].encode("utf-8").decode('unicode-escape') for c in range(len(bn_m)) ]
        # print("UTF-8 ENCODE", bn_m_en)
        optxt = ""
        try:
            
            # bn_m_en = translate.translate(bn_m)
            # print("TTT: ",bn_m_en)
            # optxt = [ engine.convert(bn_m_en[c]) for c in range(len(bn_m_en)) ]
            # print("Converted Text:", optxt)
            optxt_u = test.convertUnicodeToBijoy(bn_m)
            opttxt = test.convertBijoyToUnicode(optxt_u) #আমাদের
            # print("Converted Text:", optxt) #কায়েদা প্রতিষ্ঠার
        except TypeError:
            pass
        except IndexError:
            print("Index Error: ",bn_m)
            pass

        try:
            build_text+=f"(<p>{wrd} <font face='Nikosh'>{opttxt}</font></p>) "
            # document.add_paragraph(f"({wrd} {opttxt})")
        except NameError:
            pass
        except IndexError:
            pass

    flow_obj = Paragraph(build_text)
    build_flow_obj.append(flow_obj)
    # print(build_flow_obj)
sstr = """</body>
            </html>
        """
flow_obj = Paragraph(sstr)
build_flow_obj.append(flow_obj)
pdf.build(build_flow_obj)

# document.save('sample.docx')

endtime = time.time()
print(f"Time taken to generate PDF is: {(endtime-starttime)/60.0} minutes")


# s = u"This is an unicode string".encode('utf-8-sig')
# print s # You will see the BOM
# print s.decode('utf-8-sig')
# PythonDownload
# ï»¿This is an unicode string
# This is an unicode string
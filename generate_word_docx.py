"# -*- coding: utf-8 -*-"
import nltk
# import goslate
import pymongo
import time
import fitz  # this is pymupdf
import re
# from docx2pdf import convert
from docx import Document
from docx.shared import Pt, RGBColor


stopwords = nltk.corpus.stopwords.words('english') # take the list of enlish stopword(like the, a, an, he, she, his, him etc. )


def remove_extra_spaces(txt): 
    str = txt.split()
    new_txt=""
    for word in str:
        new_txt += word+" "
    return new_txt

def remove_stopwords(txt_tokenized):
    txt_clean = [word for word in txt_tokenized if word not in stopwords]
    return txt_clean

def remove_duplicate_word(txt, token):
    txt_clean = [word for word in txt if word not in token]
    return txt_clean

def collect_duplicate_words(txt1, txt2) :
        txt = [word for word in txt1 if word in txt2]
        return txt

# filename = r'D:\PythonProject\dictionary_project\Malgudi Days.pdf'
filename2 = r'D:\PythonProject\dictionary_project\ASURA TALE OF THE VANQUISHED The Story of Ravana.pdf'
# filename3 = r'D:\PythonProject\dictionary_project\Dharmayoddha Kalki _ Avatar of Vishnu.pdf'
# filename4 = r'D:\PythonProject\dictionary_project\Digital Electronics And Logic Design.pdf'
# filename5 = r'D:\PythonProject\dictionary_project\Catcher in the Rye.pdf'
# filename6 = r'D:\PythonProject\dictionary_project\Who Moved My Cheese__ An Amazing Way to Deal with.pdf'

client = pymongo.MongoClient("mongodb://localhost:27017/")
mydb = client["bengali_dictionary_app"]
mycol = mydb["words"]

doc = fitz.open(filename2)     # or fitz.Document(filename)
# doc.title = 'Who Moved My Cheese_'
starttime = time.time()
document = Document()

token_words = []

for page in doc:

    # print(page.getText())
    # if page.number == 4:
    #     break
    print(f"Page: {page.number}")

    text = page.getText()
    lowerText = text.lower()
    # remove_extra_spaces(text)
    words = []
    words_meanings = []
    duplicate_word_list = []
    word_list = []
    # print(stopwords[0:10])
    data = re.split("[\s.,!?:;'\"-]+",lowerText)
    # print(words)
    cln_data = remove_stopwords(data)
    words = list(dict.fromkeys(cln_data)) #remove duplicate word from the list
    
    if page.number > 2 and token_words != []:
        duplicate_word_list = collect_duplicate_words(words, token_words)
        _cln_data = remove_duplicate_word(words,duplicate_word_list)
        words = list(dict.fromkeys(_cln_data))
        token_words.extend(words)
        print(words)
    # print(words)
    else:
        token_words.extend(words)
        print(token_words)

    style = document.styles['Normal']
    font = style.font
    font.name = "Kalpurush"
    font.size = Pt(16)

    paragraph = document.add_paragraph(remove_extra_spaces(text))
    paragraph.add_run("\n\n")

    for word in words:
        wr = mycol.find({"word": word.upper()})
        
        if wr.count() != 0:
            try:
                bn_meaning = [w["bn_meanings"] for w in wr][0]
                d = word, bn_meaning
                words_meanings.append(d)
            except IndexError:
                pass
            except TypeError:
                pass
        else:
            with open(r'D:\PythonProject\dictionary_project\data\unkwon_word.txt','a') as f:
                f.write(word)
                f.write('\n')
                f.close()

    for meanings in words_meanings:
        wrd = meanings[0]
        bn_m = meanings[1]
        
        try:
            # word_list += f"({wrd} {bn_m}) "
            
            paragraph.add_run(f"({wrd} ")
            paragraph.add_run(f"{bn_m}) ").bold = True
            # font = paragraph.add_run().font
            # font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

        except TypeError:
            pass
        except IndexError:
            print("Index Error: ",bn_m)
            pass
        # print(word_list)
        # font.size = Pt(10)
        # font.name = 'Nikosh'
    # document.add_paragraph(word_list)

document.add_page_break()
document.save('sample.docx')
# convert(r"D:\PythonProject\dictionary_project\sample.docx", r"D:\PythonProject\dictionary_project\output.pdf")
endtime = time.time()
print(f"Time taken to generate PDF is: {(endtime-starttime)/60.0} minutes")

# str1 = ['good', 'mean', 'left', 'schools', 'places', 'even', 'know', 'leaving', 'hate', 
# 'care', 'sad', 'bad', 'goodby', 'leave', 'place', 'like', 'feel', 'worse', 'lucky', 
# 'sudden', 'thought', 'something', 'helped', 'make', 'getting', 'hell', 'suddenly', 
# 'remembered', 'time', 'around', 'october', 'robert', 'tichener', 'paul', 'campbell',
# 'chucking', 'football', 'front', 'academic', 'building', 'nice', 'guys', 'especially', 
# 'dinner', 'pretty', 'dark', 'kept', 'ball', 'anyway', 'darker', 'could', 'hardly', 'see', 
# 'want', 'stop', 'finally', 'teacher', 'taught', 'biology', 'mr', 'zambesi', 'stuck', 'head', 
# 'window', 'told', 'us', 'go', 'back', 'dorm', 'get', 'ready', 'chance', 'remember', 'kind', 
# 'stuff', 'need', 'one', 'least', 'soon', 'got', 'turned', 'started', 'running', 'side', 
# 'hill', 'toward', 'old', 'spencer','house', 'live', 'campus', 'lived', 'anthony', 'wayne', 
# 'avenue', 'ran', 'way', 'main', 'gate', 'waited', 'second', 'till', 'breath', 'wind', 'truth', 
# 'quite', 'heavy', 'smoker', 'thing', 'used', 'made', 'cut', 'another', 'grew', 'six', 'half', 'inches', 
# 'last', 'year', 'also', 'practically', 'b', 'came', 'goddam', 'checkups', 'healthy', 'though', 'across', 
# 'route', '204', 'icy', 'damn', 'near', 'fell', 'guess', 'felt', 'road', 'sort', 'disappearing', 'crazy', 
# 'afternoon', 'terrifically', 'cold', 'sun', 'anything', 'every', 'crossed', 'boy', 'rang', 'doorbell', 'fast', 
# 'really', 'frozen', 'ears', 'hurting', 'move', 'fingers', 'c', 'mon', 'said', 'right', 'loud', 'almost', 
# 'somebody', 'open', 'door', 'mrs', 'opened', 'maid', 'always', 'much', 'dough', 'holden', 'lovely', 
# 'come', 'dear', 'death', 'think', 'glad', 'liked', 'let', 'take', 'coat', 'hear', 'ask', 'deaf', 'hung', 
# 'hall', 'closet', 'brushed', 'hair', 'hand', 'wear', 'crew', 'frequently', 'never', 'comb', 'louder', 'fine', 
# 'closed', 'asked', 'knew', 'away', 'kicked', 'grippe', 'yet', 'behaving', 'perfect', 'room', '']

# str2 = ['2', 'room', 'around', 'seventy', 'years', 'old', 'even', 'got', 'bang', 'things', 'though', 'haif', 'assed', 
# 'way', 'course', 'know', 'sounds', 'mean', 'say', 'used', 'think', 'spencer', 'quite', 'lot', 'thought', 
# 'much', 'wondered', 'heck', 'still', 'living', 'stooped', 'terrible', 'posture', 'class', 'whenever', 'dropped', 
# 'piece', 'chalk', 'blackboard', 'guy', 'first', 'row', 'always', 'get', 'pick', 'hand', 'awful', 'opinion', 
# 'enough', 'could', 'figure', 'bad', 'instance', 'one', 'sunday', 'guys', 'hot', 'chocolate', 'showed', 'us', 
# 'beat', 'navajo', 'blanket', 'mrs', 'bought', 'indian', 'yellowstone', 'park', 'tell', 'big', 
# 'buying', 'take', 'somebody', 'hell', 'like', 'door', 'open', 'sort', 'knocked', 'anyway', 'polite', 'see', 
# 'sitting', 'leather', 'chair', 'wrapped', 'told', 'looked', 'yelled', 'caulfield', 'come', 'boy', 'yelling', 
# 'outside', 'nerves', 'sometimes', 'minute', 'went', 'sorry', 'reading', 'atlantic', 'monthly', 'pills', 'medicine', 
# 'place', 'everything', 'smelled', 'vicks', 'nose', 'drops', 'pretty', 'depressing', 'crazy', 'sick', 'people', 'made', 
# 'sad', 'ratty', 'bathrobe', 'probably', 'born', 'something', 'pajamas', 'bathrobes', 'bumpy', 'chests', 'showing', 'legs', 
# 'beaches', 'places', 'look', 'white', 'unhairy', 'hello', 'sir', 'said', 'note', 'thanks', 'written', 'asking', 'stop', 
# 'good', 'vacation', 'started', 'account', 'coming', 'back', 'seat', 'meant', 'bed', 'sat', 'grippe', 'felt', 'better', 
# 'send', 'doctor', 'chuckling', 'madman', 'finally', 'straightened', 'game', 'day', 'new', 'york', 'fencing', 'team', 'rock',
# 'getting', 'serious', 'knew', 'would', 'leaving', 'eh', 'yes', 'guess', 'going', 'nodding', 'routine', 'never', 'saw', 
# 'anybody', 'nod', 'life', 'thinking', 'nice', 'ass', 'elbow', '']

# str3 = ['', 'may', 'ask', 'oh', 'well', 'long', 'story', 'sir', 'mean', 'pretty', 'complicated', 'feel', 'like', 'going', 'whole', 'thing', 'understood', 'anyway', 'alley', 'one', 'biggest', 'reasons', 'left', 'elkton', 'hills', 'surrounded', 'phonies', 'coming', 'goddam', 'window', 'instance', 'headmaster', 'mr', 'haas', 'phoniest', 'bastard', 'ever', 'met', 'life', 'ten', 'times', 'worse', 'old', 'thurmer', 'sundays', 'went', 'around', 'shaking', 'hands', 'everybody', 'parents', 'drove', 'school', 'charming', 'hell', 'except', 'boy', 'little', 'funny', 'looking', 'seen', 'way', 'roommate', 'mother', 'sort', 'fat', 'corny', 'something', 'somebody', 'father', 'guys', 'wear', 'suits', 'big', 'shoulders', 'black', 'white', 'shoes', 'hans', 'would', 'shake', 'give', 'phony', 'smile', 'go', 'talk', 'maybe', 'half', 'hour', 'else', 'stand', 'stuff', 'drives', 'crazy', 'makes', 'depressed', 'hated', 'spencer', 'asked', 'hear', 'thinking', 'said', 'particular', 'qualms', 'leaving', 'pencey', 'right', 'sure', 'many', 'yet', 'guess', 'really', 'hit', 'takes', 'things', 'home', 'wednesday', 'moron', 'absolutely', 'concern', 'future', 'thought', 'minute', 'much', 'late', 'hearing', 'say', 'made', 'sound', 'dead', 'depressing', 'put', 'sense', 'head', 'trying', 'help', 'could', 
# 'see', 'opposite', 'sides', 'ot', 'pole', 'know', 'thanks', 'lot', 'kidding', 'appreciate', 'got', 'bed', 'sat', 'another', 'minutes', 'save', 'though', 'get', 
# 'quite', 'bit', 'equipment', 'gym', 'take', 'looked', 'started', 'nodding', 'serious', 'look', 'face', 'felt', 'sorry', 'sudden', 'hang', 'longer', 'kept', 'missing', 'whenever', 'chucked', 'sad', 'bathrobe', 'chest', 'showing', 'grippy', 'smell', 'vicks', 'nose', 'drops', 'place', 'worry', 'phase', 'goes', 'phases', 'hate', 'answers', 'please', 'hand', 'shoulder', 'okay']

# # ['left', 'schools', 'hate', 'care', 'goodby', 'leave', 'feel', 'worse', 'lucky', 'sudden', 'helped', 'make', 'suddenly', 
# # 'remembered', 'time', 'october', 'robert', 'tichener', 'paul', 'campbell', 'chucking', 'football', 'front', 'academic', 
# # 'building', 'especially', 'dinner', 'dark', 'kept', 'ball', 'darker', 'hardly', 'want', 'teacher', 'taught', 'biology', 
# # 'mr', 'zambesi', 'stuck', 'head', 'window', 'go', 'dorm', 'ready', 'chance', 'remember', 'kind', 'stuff', 'need', 'least', 
# # 'soon', 'turned', 'running', 'side', 'hill', 'toward', 'house', 'live', 'campus', 'lived', 'anthony', 'wayne', 'avenue',
# # 'ran', 'main', 'gate', 'waited', 'second', 'till', 'breath', 'wind', 'truth', 'heavy', 'smoker', 'thing', 'cut', 'another', 
# # 'grew', 'six', 'half', 'inches', 'last', 'year', 'also', 'practically', 'b', 'came', 'goddam', 'checkups', 'healthy', 
# # 'across', 'route', '204', 'icy', 'damn', 'near', 'fell', 'road', 'disappearing', 'afternoon', 'terrifically', 'cold', 
# # 'sun', 'anything', 'every', 'crossed', 'rang', 'doorbell', 'fast', 'really', 'frozen', 'ears', 'hurting', 'move', 
# # 'fingers', 'c', 'mon', 'right', 'loud', 'almost', 'opened', 'maid', 'dough', 'holden', 'lovely', 'dear', 'death', 'glad', 
# # 'liked', 'let', 'coat', 'hear', 'ask', 'deaf', 'hung', 'hall', 'closet', 'brushed', 'hair', 'wear', 'crew', 'frequently', 
# # 'comb', 'louder', 'fine', 'closed', 'asked', 'away', 'kicked', 'yet', 'behaving', 'perfect']
# str4 = ['', 'like', 'cup', 'hot', 'chocolate', 'go', 'mrs', 'spencer', 'would', 'really', 'thing', 'get', 
# 'going', 'right', 'gym', 'thanks', 'though', 'lot', 'sir', 'shook', 'hands', 'crap', 'made', 'feel', 'sad', 'hell',
# 'drop', 'line', 'take', 'care', 'grippe', 'good', 'boy', 'shut', 'door', 'started', 'back', 'living', 'room', 'yelled', 
# 'something', 'exactly', 'hear', 'pretty', 'sure', 'luck', 'hope', 'never', 'yell', 'anybody', 'sounds', 'terrible', 
# 'think', '3', 'terrific', 'liar', 'ever', 'saw', 'life', 'awful', 'way', 'store', 'buy', 'magazine', 'even', 'somebody', 
# 'asks', 'liable', 'say', 'opera', 'told', 'old', 'equipment', 'stuff', 'sheer', 'lie', 'keep', 'goddam', 'lived', 
# 'pencey', 'ossenburger', 'memorial', 'wing', 'new', 'dorms', 'juniors', 'seniors', 'junior', 'roommate', 'senior', 
# 'named', 'guy', 'went', 'pot', 'dough', 'undertaking', 'business', 'got', 'parlors', 'country', 'could', 'members', 
# 'family', 'buried', 'five', 'bucks', 'apiece', 'see', 'probably', 'shoves', 'sack', 'dumps', 'river', 'anyway', 'gave', 
# 'pile', 'alter', 'first', 'football', 'game', 'year', 'came', 'school', 'big', 'cadillac', 'stand', 'grandstand', 
# 'give', 'locomotive', 'cheer', 'next', 'morning', 'chapel', 'speech', 'lasted', 'ten', 'hours', 'fifty', 'corny', 
# 'jokes', 'show', 'us', 'regular', 'deal', 'telling', 'ashamed', 'kind', 'trouble', 'knees', 'pray', 'god', 'always', 
# 'talk', 'wherever', 'ought', 'jesus', 'buddy', 'said', 'talked', 'time', 'driving', 'car', 'killed', 'phony', 
# 'bastard', 'shifting', 'gear', 'asking', 'send', 'stiffs', 'part', 'middle', 'swell', 'shot', 'sudden', 'sitting', 
# 'row', 'front', 'edgar', 'marsalla', 'laid', 'fart', 'crude', 'also', 'quite', 'amusing', 'damn', 'near', 'blew', 
# 'roof', 'hardly', 'laughed', 'loud', 'thurmer', 'headmaster', 'rostrum', 'tell', 'heard', 'sore', 'anything', 'night', 
# 'compulsory', 'study', 'hall', 'academic', 'building', 'created']

# str5 = ['cup', 'drop', 'line', 'shut', 'luck', 'hope', 'yell', '3', 'liar', 'store', 'buy', 'magazine', 'asks', 'liable', 'opera', 'sheer', 'lie', 'ossenburger', 'memorial', 'wing', 'dorms', 'juniors', 'junior', 'senior', 'named', 'pot', 'undertaking', 'business', 'parlors', 'country', 'members', 'family', 'buried', 'shoves', 'sack', 'dumps', 'river', 'alter', 'cadillac', 'locomotive', 'cheer', 'chapel', 'speech', 'lasted', 'fifty', 'jokes', 'show', 'ashamed', 'trouble', 'knees', 'pray', 'god', 'wherever', 'ought', 'jesus', 'buddy', 'talked', 'driving', 'car', 'shifting', 'gear', 'stiffs', 'part', 'middle', 'swell', 'edgar', 'marsalla', 'laid', 'fart', 'crude', 'amusing', 'blew', 'roof', 'laughed', 'rostrum', 'sore', 'compulsory', 'study', 'created']
# # print(remove_duplicate_word(str1,str3))
# print(collect_duplicate_words(str1,str5))